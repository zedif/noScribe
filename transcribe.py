#!/usr/bin/env python

import os
import platform
import yaml
import locale
import appdirs
from subprocess import run, Popen, PIPE, STDOUT
if platform.system() == 'Windows':
    from subprocess import STARTUPINFO, STARTF_USESHOWWINDOW
from docx import Document
import docx
import re

from threading import Thread
from queue import Queue, Empty
import datetime
from pathlib import Path

import pickle
import argparse
from time import sleep

if platform.system() == 'Windows':
    import cpufeature
if platform.system() in ["Darwin", "Linux"]: # = macOS or Linux
    import shlex


if platform.system() in ["Darwin", "Linux"]: # = macOS or Linux
    bundle_dir = os.path.abspath(os.path.dirname(__file__))

app_version = '0.3'

import i18n
from i18n import t
i18n.set('filename_format', '{locale}.{format}')
i18n.load_path.append('./trans')
try:
    app_locale = locale.getdefaultlocale()[0][0:2]
except:
    app_locale = 'en'
i18n.set('fallback', 'en')
i18n.set('locale', app_locale)

# Check CPU capabilities and select the right version of whisper
if platform.system() == 'Windows':
    if cpufeature.CPUFeature["AVX2"] == True and cpufeature.CPUFeature["OS_AVX"] == True:
        whisper_path = "./whisper_avx2"
    else:
        whisper_path = "./whisper_sse2"
elif platform.system() == "Darwin": # = macOS
    if platform.machine() == "arm64":
        whisper_path = "./whisper_mac"
    elif platform.machine() == "x86_64":
        raise Exception('Platform not supported yet.')
    else:
        raise Exception('Could not detect Apple architecture.')
elif platform.system() == "Linux":
    if platform.machine() == "x86_64":
        whisper_path = "./whisper_linux"
    else:
        raise Exception('Platform not supported yet.')
else:
    raise Exception('Platform not supported yet.')

WHISPER_MODELPATH = './models/'
DEFAULT_WHISPER_MODEL = os.path.abspath(WHISPER_MODELPATH + 'ggml-base.bin')

speaker_detection = 'auto'
start = 0

# timestamp regex
timestamp_re = re.compile('\[\d\d:\d\d:\d\d.\d\d\d --> \d\d:\d\d:\d\d.\d\d\d\]')

# Helper functions

def millisec(timeStr): # convert 'hh:mm:ss' string to milliseconds
    try:
        spl = timeStr.split(':')
        s = (int)((int(spl[0]) * 60 * 60 + int(spl[1]) * 60 + float(spl[2]) )* 1000)
        return s
    except:
        raise Exception(t('err_invalid_time_string', time = timeStr))

def iter_except(function, exception):
        # Works like builtin 2-argument `iter()`, but stops on `exception`.
        try:
            while True:
                yield function()
        except exception:
            return
        
def docx_add_bookmark(first_run, last_run, bookmark_name, bookmark_id):
    # adds a bookmark including the two runs and everything inbetween 
    # bookmark_id must be unique
    start = docx.oxml.shared.OxmlElement('w:bookmarkStart')
    start.set(docx.oxml.ns.qn('w:id'), str(bookmark_id))
    start.set(docx.oxml.ns.qn('w:name'), bookmark_name)
    first_run._r.append(start)

    end = docx.oxml.shared.OxmlElement('w:bookmarkEnd')
    end.set(docx.oxml.ns.qn('w:id'), str(bookmark_id))
    end.set(docx.oxml.ns.qn('w:name'), bookmark_name)
    last_run._r.append(end)

def cli():

    parser = argparse.ArgumentParser(description=t('app_header'))
    parser.add_argument("-w", "--wav-input", metavar="file_name", help="Input wave file", required=True)
    parser.add_argument("-d", "--diarization-input", metavar="file_name", help="Input diarization file (pickle file)", required=True)
    parser.add_argument("-o", "--output", metavar="file_name", help="Output transcript file (.docm)", required=True)
    parser.add_argument("-m", "--model", metavar="file_name", default="./models/ggml-base.bin",
        help="path to whisper model, e.g. ./models/ggml-base.bin")
    parser.add_argument("-l", "--language", default='auto', help="spoken language ('en' for English, 'de' for German, 'auto' for auto-detect)")
    parser.add_argument("-t", "--threads", default="4", help="Number of parallel threads to use (default: 4)")
    parser.add_argument(
        "--auto-save",
        action="store_true",
        help="Enable auto-save",
    )
    parser.add_argument("--max-len", default="30", metavar='N', help="whisper.cpp flag: maximum segment length in characters")
    args = parser.parse_args()

    transcribe(
        wav_audio_file=args.wav_input,
        diarization_file=args.diarization_input,
        transcript_file=args.output,
        language=args.language,
        whisper_model=os.path.abspath(args.model),
        auto_save=args.auto_save,
        whisper_options="--max-len " + args.max_len,
        number_of_threads=args.threads
    )


def __load_pickle(file_path):
    try:
        with open(file_path, 'rb') as file:
            return pickle.load(file)
    except FileNotFoundError:
        print(f"Error: The diarization file '{file_path}' was not found.")
        return None
    except pickle.UnpicklingError:
        print(f"Error: Failed to unpickle the diarization from '{file_path}'.")
        return None
    except Exception as e:
        print(f"An unexpected error occurred: {e}")
        return None


def reader_thread(process, q):
    try:
        with process.stdout as pipe:
            for line in iter(pipe.readline, b''):
                q.put(line)
    finally:
        q.put(None)


def transcribe(wav_audio_file, diarization_file, transcript_file, language='auto', whisper_model=DEFAULT_WHISPER_MODEL,
               whisper_options="--max-len 30", number_of_threads="4", whisper_extra_commands='', auto_save=True):

    print(t('welcome_message'))
    print(t('welcome_credits', v=app_version))
    print('https://github.com/kaixxx/noScribe')

    # TODO: audio_file should contain the original compressed audio file (.mp3, .m4a, etc.) that will be referenced
    #       in the Word transcript, but this routine doesn't know it yet.
    audio_file = wav_audio_file

    proc_start_time = datetime.datetime.now()
    transcription_quality = 'fast'

    print(f'Saving transcript to: {transcript_file}')

    try:
        # log CPU capabilities
        if platform.system() == 'Windows':
            print("=== CPU FEATURES ===")
            for key, value in cpufeature.CPUFeature.items():
                print('    {:24}: {}'.format(key, value))

        try:

            def overlap_len(ss_start, ss_end, ts_start, ts_end):
                # ss...: speaker section start and end in milliseconds (from pyannote)
                # ts...: transcript section start and end (from whisper.cpp)
                if ts_end < ss_start: # no overlap, ts is before ss
                    return -1
                elif ts_start > ss_end: # no overlap, ts is after ss
                    return 0
                else: # ss & ts have overlap
                    if ts_start > ss_start: # ts starts after ss
                        overlap_start = ts_start
                    else:
                        overlap_start = ss_start
                    if ts_end > ss_end: # ts ends after ss
                        overlap_end = ss_end
                    else:
                        overlap_end = ts_end
                    return overlap_end - overlap_start + 1

            def find_speaker(diarization, transcript_start, transcript_end):
                # Looks for the segment in diarization that has the most overlap with section_start-end.
                # Returns the speaker name if found, an empty string otherwise
                spkr = ''
                overlap = 0

                for segment, _, label in diarization.itertracks(yield_label=True):
                    t = overlap_len(int(segment.start * 1000), int((segment.start + segment.duration) * 1000), transcript_start, transcript_end)
                    if t == -1: # we are already after transcript_end
                        break
                    elif t > overlap:
                        overlap = t
                        spkr = f'S{label[8:]}' # shorten the label: "SPEAKER_01" > "S01"
                return spkr

            #-------------------------------------------------------
            # 2) Load in corresponding diarization
            diarization = __load_pickle(diarization_file)
            print(f'Loaded diarization file: {diarization_file}')

            #-------------------------------------------------------
            # 3) Transcribe with whisper.cpp

            print(t('start_transcription'))
            print(t('loading_whisper'))

            command = f'{whisper_path}/main --model {whisper_model} --language {language} {whisper_options} --threads {number_of_threads} --print-colors --print-progress --file "{wav_audio_file}" {whisper_extra_commands}'
            if platform.system() in ["Darwin", "Linux"]: # = macOS or Linux
                command = shlex.split(command)
            print(f'Whisper command: {command}')

            # prepare transcript docm
            d = Document('transcriptTempl.docm')
            d.core_properties.author = f'noScribe vers. {app_version}'
            d.core_properties.comments = audio_file

            # header
            p = d.paragraphs[0]
            p.text = Path(audio_file).stem # use the name of the audio file (without extension) as the title
            p.style = 'noScribe_header'

            p = d.add_paragraph(t('doc_header', version=app_version), style='noScribe_subheader')
            p = d.add_paragraph(t('doc_header_audio', file=audio_file), style='noScribe_subheader')

            p = d.add_paragraph()
            speaker = ''
            bookmark_id = 0
            last_auto_save = datetime.datetime.now()

            try:
                if platform.system() == 'Windows':
                    startupinfo = STARTUPINFO()
                    startupinfo.dwFlags |= STARTF_USESHOWWINDOW
                    process = Popen(command, stdout=PIPE, stderr=STDOUT, startupinfo=startupinfo)
                elif platform.system() in ["Darwin", "Linux"]:
                    process = Popen(command, stdout=PIPE, stderr=STDOUT)
                # Run whisper.cpp main.exe without blocking the GUI:
                # Source: https://stackoverflow.com/questions/12057794/python-using-popen-poll-on-background-process
                # launch thread to read the subprocess output
                #   (put the subprocess output into the queue in a background thread,
                #    get output from the queue in the GUI thread.
                #    Output chain: process.readline -> queue -> GUI)
                q = Queue(maxsize=1024)  # limit output buffering (may stall subprocess)
                th = Thread(target=reader_thread, args=[process, q])
                th.daemon = True # close pipe if GUI process exits
                th.start()

                # TODO: capture CTRL-C -> cancel = True
                cancel = False
                while process.poll() == None: # process is running

                    # poll only every 1/10th of a second for the while loop not to max out one core.
                    sleep(0.1)

                    # check for user cancelation
                    if cancel == True:
                        if auto_save == True:
                            d.save(transcript_file)
                            print()
                            print(t('transcription_saved', file=transcript_file))
                            raise Exception(t('err_user_cancelation'))
                        else:
                            raise Exception(t('err_user_cancelation'))

                    # process lines from the queue
                    for line in iter_except(q.get_nowait, Empty):
                        if line is None:
                            break
                        else:
                            line = str(line.decode("utf-8", errors='ignore')) # convert to regular string

                            # check if we have a transcript line from stdout or a line from stdterr
                            if timestamp_re.match(line) != None:
                                # found a timestamp, must be a transcript

                                line = line.replace('\n', '') # remove line breaks
                                line = line.replace('\r', '') # remove carriage return

                                # get time of the segment in milliseconds
                                #[00:00:00.000 --> 00:00:05.760]
                                d_start = line[1:13]
                                d_end = line[18:30]
                                d_start = millisec(d_start)
                                d_end = millisec(d_end)

                                line = line[33:] # discard timestamp
                                line = line.lstrip() # discard leading spaces

                                # write text to the doc
                                # diarization (speaker detection)?
                                if speaker_detection == 'auto':
                                    spkr = find_speaker(diarization, d_start, d_end)
                                    if (speaker != spkr) & (spkr != ''):
                                        speaker = spkr
                                        print()
                                        p = d.add_paragraph()
                                        line = f'{speaker}: {line}'

                                first_run = p.add_run() # empty run for start_bookmark
                                # check for confidence level markers (colors)
                                if line.find('\u001B[38;5;') > -1:
                                    line_segments = line.split('\u001B[38;5;')
                                    cl_markers = {'196m': 1, '202m': 2, '208m': 3, '214m': 4, '220m': 5, '226m': 6, '190m': 7, '154m': 8, '118m': 9, '82m': 10}
                                    for s in line_segments:
                                        if s == '':
                                            continue # skip empty segments
                                        # extract confidence level marker, get the level from cl_markers:
                                        cl_marker_end = s.find('m')
                                        if cl_marker_end in [2,3]: # only possible positions
                                            cl_marker = s[0:cl_marker_end + 1]
                                            if cl_marker in cl_markers:
                                                cl_level = cl_markers[cl_marker]
                                            else: # invalid marker
                                                cl_level = 0
                                        else: # marker not found
                                            cl_level = 0
                                        # add segments to doc
                                        s = s[cl_marker_end + 1:] # delete cl_marker
                                        s = s.replace('\u001B[0m', '') # delete the closing cl mark
                                        r = p.add_run()
                                        r.text = s
                                        # Mark confidence level with a character based style,'noScribe_cl[1-10]'
                                        # This way, we can color-mark specific levels later in Word.
                                        if cl_level > 0:
                                            r.style = d.styles[f'noScribe_cl{cl_level}']
                                        print(s, end='', sep='')
                                else: # no marker in line
                                        r = p.add_run()
                                        r.text = line
                                        print(line, end='', sep='')

                                # Create bookmark with audio timestamps start to end.
                                # This way, we can jump to the according audio position and play it later in Word.
                                bookmark_id = bookmark_id + 1
                                last_run = p.add_run()
                                # if we skipped a part at the beginning of the audio we have to add this here again, otherwise the timestaps will not match the original audio:
                                orig_audio_start = start + d_start
                                orig_audio_end = start + d_end
                                docx_add_bookmark(first_run, last_run, f'ts_{orig_audio_start}_{orig_audio_end}', bookmark_id)

                                # auto save
                                if auto_save == True:
                                    if (datetime.datetime.now() - last_auto_save).total_seconds() > 20:
                                        d.save(transcript_file)


                            else: # must be line from stderr
                                print(line)

                print(f'trying to safe to {transcript_file}:')
                d.save(transcript_file)
                print()
                print()
                print(t('transcription_finished'))

                print(t('transcription_saved'), transcript_file)

                # log duration of the whole process in minutes
                proc_time = datetime.datetime.now() - proc_start_time
                print(t('transcription_time', duration=int(proc_time.total_seconds() / 60)))

                if process.poll() > 0:
                    raise Exception(t('err_whisper_main', e=process.poll()))

            except Exception as e:
                print()
                print(t('err_transcription'), 'error')
                print(e, 'error')
                return

            finally:
                process.kill() # exit subprocess (zombie!)

        except Exception as e:
            print(f"An unexpected error occurred: {e}")


    except Exception as e:
        print(t('err_options'), 'error')
        print(e, 'error')
        return


if __name__ == "__main__":
    cli()
